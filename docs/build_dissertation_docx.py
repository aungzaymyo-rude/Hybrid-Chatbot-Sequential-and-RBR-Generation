from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS_DIR = ROOT / "docs"
FIG_DIR = DOCS_DIR / "figures"
MD_PATH = DOCS_DIR / "dissertation_draft.md"
OUT_PATH = DOCS_DIR / "dissertation_final.docx"


FIGURE_MAP = {
    "Figure 3.1": FIG_DIR / "figure_3_1_hybrid_architecture.png",
    "Figure 3.2": FIG_DIR / "figure_3_2_mlops_lifecycle.png",
    "Figure 3.3": FIG_DIR / "figure_3_3_chat_ui.png",
    "Figure 4.1": FIG_DIR / "figure_4_1_metric_comparison.png",
    "Figure 4.2": FIG_DIR / "figure_4_2_training_curves.png",
    "Figure 4.3": FIG_DIR / "figure_4_3_dataset_and_splits.png",
    "Figure 4.4": FIG_DIR / "figure_4_4_testing_summary.png",
    "Figure 4.5": FIG_DIR / "figure_4_5_admin_overview.png",
    "Figure 4.6": FIG_DIR / "figure_4_7_admin_trace.png",
}


def set_default_font(style, name: str = "Arial", size: int = 11) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_document_layout(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    set_default_font(doc.styles["Normal"])
    set_default_font(doc.styles["Title"], size=20)
    set_default_font(doc.styles["Heading 1"], size=16)
    set_default_font(doc.styles["Heading 2"], size=13)
    if "Heading 3" in doc.styles:
        set_default_font(doc.styles["Heading 3"], size=11)


def set_para_spacing(paragraph, line_15: bool = True) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(6)
    fmt.space_before = Pt(0)
    if line_15:
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        fmt.line_spacing = 1.0


def add_page_field(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def restart_page_numbering(section) -> None:
    sect_pr = section._sectPr
    pg_num_type = sect_pr.find(qn("w:pgNumType"))
    if pg_num_type is None:
        pg_num_type = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num_type)
    pg_num_type.set(qn("w:start"), "1")


def add_toc(paragraph) -> None:
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), 'TOC \\o "1-3" \\h \\z \\u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Right-click and update field in Word to generate the table of contents."
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def clean_inline_markdown(text: str) -> str:
    text = text.replace("**", "")
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    return text


def add_body_paragraph(doc: Document, text: str, style: str | None = None, single_space: bool = False):
    p = doc.add_paragraph(style=style)
    p.add_run(clean_inline_markdown(text))
    set_para_spacing(p, line_15=not single_space)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    set_para_spacing(p, line_15=False)
    return p


def add_figure(doc: Document, image_path: Path, caption: str, width: Inches = Inches(6.3)):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=width)
    set_para_spacing(p, line_15=False)
    add_caption(doc, caption)


def add_metrics_table(doc: Document):
    add_caption(doc, "Table 4.1 Final test-set metrics for the general and report model profiles")
    table = doc.add_table(rows=1, cols=6)
    table.style = "Table Grid"
    headers = ["Model", "Accuracy", "Macro F1", "Weighted F1", "Macro Precision", "Macro Recall"]
    for cell, header in zip(table.rows[0].cells, headers):
        cell.text = header
    rows = [
        ["General", "0.8928", "0.8649", "0.8914", "0.8822", "0.8617"],
        ["Report", "0.9040", "0.8604", "0.8991", "0.8660", "0.8592"],
    ]
    for row in rows:
        cells = table.add_row().cells
        for cell, value in zip(cells, row):
            cell.text = value
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_para_spacing(p, line_15=False)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10.5)


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Design and Development of a Hybrid Chatbot for Medical Haematology using Sequential Models for Intent Classification and Retrieval-Based Response Generation")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(20)

    for line in [
        "",
        "Dissertation Draft for Final Year Project",
        "Programme: [Insert Degree Title]",
        "Student Name: [Insert Name]",
        "Student ID: [Insert ID]",
        "Supervisor: [Insert Supervisor Name]",
        "Submission Year: 2026",
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(line).font.name = "Arial"
        set_para_spacing(para, line_15=False)

    doc.add_page_break()


def add_front_matter(doc: Document):
    add_body_paragraph(doc, "Abstract", style="Heading 1", single_space=True)
    md_text = MD_PATH.read_text(encoding="utf-8")
    abstract = md_text.split("## Abstract", 1)[1].split("## Chapter 1", 1)[0].strip().split("\n\n")
    for para in abstract:
        add_body_paragraph(doc, para.strip(), single_space=True)

    doc.add_page_break()
    add_body_paragraph(doc, "Acknowledgements", style="Heading 1", single_space=True)
    add_body_paragraph(
        doc,
        "Support provided by the project supervisor, peer reviewers, and testers was acknowledged. Technical feedback received during iterative development was incorporated into the final artefact and documentation.",
        single_space=True,
    )

    doc.add_page_break()
    add_body_paragraph(doc, "Table of Contents", style="Heading 1", single_space=True)
    toc_para = doc.add_paragraph()
    add_toc(toc_para)
    set_para_spacing(toc_para, line_15=False)


def insert_placeholder_asset(doc: Document, line: str):
    if "Figure 3.1" in line:
        add_figure(doc, FIGURE_MAP["Figure 3.1"], "Figure 3.1 Hybrid haematology chatbot architecture", Inches(6.4))
    elif "Table 4.1" in line:
        add_metrics_table(doc)
    elif "Figure 4.1. Automated test execution summary" in line:
        add_figure(doc, FIGURE_MAP["Figure 4.4"], "Figure 4.4 Automated test execution summary", Inches(5.8))
    else:
        add_body_paragraph(doc, clean_inline_markdown(line))


def add_main_body(doc: Document):
    md_text = MD_PATH.read_text(encoding="utf-8")
    main = md_text.split("## Chapter 1", 1)[1].split("## References", 1)[0]
    lines = ["## Chapter 1" + main]

    pending_paragraph: list[str] = []

    def flush_paragraph():
        nonlocal pending_paragraph
        if pending_paragraph:
            text = " ".join(pending_paragraph).strip()
            if text:
                add_body_paragraph(doc, text)
            pending_paragraph = []

    for raw_line in lines[0].splitlines():
        line = raw_line.rstrip()
        if not line.strip():
            flush_paragraph()
            continue
        if line.startswith("## "):
            flush_paragraph()
            add_body_paragraph(doc, line[3:], style="Heading 1")
            chapter_title = line[3:]
            if chapter_title == "Chapter 3: Design and Development":
                add_figure(doc, FIGURE_MAP["Figure 3.1"], "Figure 3.1 Hybrid haematology chatbot architecture", Inches(6.4))
                add_figure(doc, FIGURE_MAP["Figure 3.2"], "Figure 3.2 MLOps lifecycle for the implemented chatbot", Inches(6.2))
                add_figure(doc, FIGURE_MAP["Figure 3.3"], "Figure 3.3 Implemented chat user interface", Inches(6.2))
            if chapter_title == "Chapter 4: Result Analysis and Evaluation":
                add_metrics_table(doc)
                add_figure(doc, FIGURE_MAP["Figure 4.1"], "Figure 4.1 Test-set metric comparison across model profiles", Inches(6.1))
                add_figure(doc, FIGURE_MAP["Figure 4.2"], "Figure 4.2 Training and validation behaviour of the two BiLSTM models", Inches(6.3))
                add_figure(doc, FIGURE_MAP["Figure 4.3"], "Figure 4.3 Dataset footprint and fixed split allocation", Inches(6.2))
                add_figure(doc, FIGURE_MAP["Figure 4.4"], "Figure 4.4 Automated test execution summary", Inches(5.8))
                add_figure(doc, FIGURE_MAP["Figure 4.5"], "Figure 4.5 Admin overview dashboard for operational monitoring", Inches(6.3))
                add_figure(doc, FIGURE_MAP["Figure 4.6"], "Figure 4.6 Inference trace for the phrase “What is aPTT?”", Inches(6.3))
            continue
        if line.startswith("### "):
            flush_paragraph()
            add_body_paragraph(doc, line[4:], style="Heading 2")
            continue
        if re.match(r"^\d+\.\s", line):
            flush_paragraph()
            add_body_paragraph(doc, re.sub(r"^\d+\.\s", "", line), style="List Number")
            continue
        if line.startswith("- "):
            flush_paragraph()
            add_body_paragraph(doc, line[2:], style="List Bullet")
            continue
        if line.startswith("[Insert "):
            flush_paragraph()
            insert_placeholder_asset(doc, line)
            continue
        pending_paragraph.append(line)
    flush_paragraph()


def add_references_and_appendices(doc: Document):
    md_text = MD_PATH.read_text(encoding="utf-8")
    references = md_text.split("## References", 1)[1].split("## Appendices Guidance", 1)[0].strip().split("\n\n")
    appendices = md_text.split("## Appendices Guidance (for the final Word version)", 1)[1].strip().splitlines()

    add_body_paragraph(doc, "References", style="Heading 1")
    for ref in references:
        p = doc.add_paragraph()
        p.add_run(clean_inline_markdown(ref.strip()))
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        set_para_spacing(p)

    add_body_paragraph(doc, "Appendices", style="Heading 1")
    for line in appendices:
        if not line.strip():
            continue
        if line.startswith("- "):
            add_body_paragraph(doc, line[2:], style="List Bullet")
        else:
            add_body_paragraph(doc, clean_inline_markdown(line))


def build_document():
    doc = Document()
    set_document_layout(doc)
    add_title_page(doc)
    add_front_matter(doc)

    main_section = doc.add_section(WD_SECTION.NEW_PAGE)
    main_section.left_margin = Cm(3)
    main_section.right_margin = Cm(2)
    main_section.top_margin = Cm(2.5)
    main_section.bottom_margin = Cm(2.5)
    restart_page_numbering(main_section)
    footer = main_section.footer
    footer.is_linked_to_previous = False
    footer_para = footer.paragraphs[0]
    add_page_field(footer_para)

    add_main_body(doc)
    add_references_and_appendices(doc)
    doc.save(OUT_PATH)
    print(f"Saved {OUT_PATH}")


if __name__ == "__main__":
    build_document()
