# -*- coding: utf-8 -*-
from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
FIG = DOCS / "figures"
OUT_MD = DOCS / "testing_documentation.md"
OUT_DOCX = DOCS / "testing_documentation.docx"


@dataclass(frozen=True)
class TestCase:
    case_id: str
    category: str
    component: str
    objective: str
    task: str
    expected: str
    actual: str
    status: str
    evidence: str


TEST_CASES = [
    TestCase("UT-01", "Unit", "Preprocessing", "Normalize text for downstream inference.", "Input mixed-case query with punctuation.", "Normalized lowercase phrase is produced consistently.", "Passed in `test_preprocessing.py`.", "Pass", "Automated"),
    TestCase("UT-02", "Unit", "Predictor rule layer", "Detect greeting phrases before model inference.", "Submit `hi` to the predictor.", "Intent resolves to `greeting` with rule priority.", "Passed in `test_predictor_rules.py` and API smoke checks.", "Pass", "Automated + smoke"),
    TestCase("UT-03", "Unit", "Predictor rule layer", "Detect social small-talk phrases.", "Submit `how are u` to the predictor.", "Intent resolves to `small_talk` instead of `incomplete_query`.", "Passed in `test_predictor_rules.py`.", "Pass", "Automated"),
    TestCase("UT-04", "Unit", "Safety rule layer", "Block unsafe treatment language.", "Submit phrase containing `prescribe` or `inject`.", "Intent resolves to `unsafe_medical_request`.", "Passed in `test_predictor_rules.py` and route tests.", "Pass", "Automated"),
    TestCase("UT-05", "Unit", "Entity detection", "Recognize MCV as an RBC-report entity.", "Submit `What is MCV?`.", "Entity rule maps to canonical MCV question.", "Passed in `test_entity_detection.py`.", "Pass", "Automated"),
    TestCase("UT-06", "Unit", "Entity detection", "Recognize aPTT as coagulation entity.", "Submit `What is aPTT?`.", "Entity rule maps to coagulation knowledge.", "Passed via routing tests and smoke checks.", "Pass", "Automated + smoke"),
    TestCase("UT-07", "Unit", "Routing layer", "Prioritize direct sample-collection retrieval.", "Submit `Which tube is used for CBC?` with `sample_collection` intent.", "EDTA/lavender-top answer is returned instead of generic CBC info.", "Passed in `test_routing_engine.py`.", "Pass", "Automated"),
    TestCase("UT-08", "Unit", "Routing layer", "Return CBC sample rejection criteria.", "Submit `What are the rejection criteria for a CBC sample?`.", "Rejection-criteria answer mentions wrong tube / clotted sample style conditions.", "Passed in `test_routing_engine.py`.", "Pass", "Automated"),
    TestCase("UT-09", "Unit", "Routing layer", "Return CBC stability guidance.", "Submit `How long is a CBC sample stable?`.", "Stability answer references validated limits and SOP.", "Passed in `test_routing_engine.py`.", "Pass", "Automated"),
    TestCase("UT-10", "Unit", "Model advisory", "Auto-switch report questions from general to report model.", "Submit `How do I read this CBC report?` while `general` is selected.", "Model advisory recommends / routes to `report`.", "Passed in `test_model_advisory.py` and UI smoke checks.", "Pass", "Automated + smoke"),
    TestCase("UT-11", "Unit", "Model advisory", "Auto-switch workflow questions from report to general model.", "Submit `Which tube is used for coagulation tests?` while `report` is selected.", "Model advisory recommends / routes to `general`.", "Passed in `test_model_advisory.py` and UI smoke checks.", "Pass", "Automated + smoke"),
    TestCase("UT-12", "Unit", "Report analysis", "Interpret numeric WBC results.", "Submit `WBC is 13.7`.", "Intent resolves to `report_numeric_result_analysis` and response marks high range.", "Passed in `test_report_analysis.py` and manual runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-13", "Unit", "Report analysis", "Interpret printed report flag text.", "Submit `My report shows anemia`.", "Intent resolves to `report_flag_result_analysis` with bounded non-diagnostic explanation.", "Passed in `test_report_analysis.py` and runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-14", "Unit", "Report analysis", "Apply adult age context without exposing internal note.", "Submit `in my report WBC is 13.37 age is 51`.", "Adult range is applied, but no internal age-note sentence is displayed.", "Verified after age-aware retraining and response cleanup.", "Pass", "Manual + regression"),
    TestCase("UT-15", "Unit", "Report analysis", "Apply pediatric age context.", "Submit `WBC is 13 age is 10`.", "Pediatric range is applied and result is interpreted within pediatric criteria.", "Passed in `test_report_analysis.py` and runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-16", "Unit", "Report analysis", "Apply sex-specific HGB range.", "Submit `Adult male HGB is 13.0`.", "Male reference interval is used for HGB.", "Passed in `test_report_analysis.py` and runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-17", "Unit", "Report analysis", "Apply sex-specific HCT range.", "Submit `Female HCT is 45`.", "Female reference interval is used for HCT.", "Passed in `test_report_analysis.py` and runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-18", "Unit", "Report analysis", "Prevent age-only phrases from being treated as report analysis.", "Submit `age is 43`.", "Intent resolves to `incomplete_query`.", "Passed in `test_report_predictor.py` and runtime verification.", "Pass", "Automated + manual"),
    TestCase("UT-19", "Unit", "API layer", "Return stable chat response schema.", "POST `/chat` with `Hello`.", "JSON response includes `intent`, `response`, `lang`, and model metadata.", "Passed in `test_api.py`.", "Pass", "Automated"),
    TestCase("UT-20", "Unit", "API layer", "Expose report-analysis preview rows.", "GET `/admin/api/report-analysis-preview`.", "Preview returns report-analysis retraining candidates in JSON.", "Passed in `test_api.py`.", "Pass", "Automated"),
    TestCase("UT-21", "Unit", "Admin trace API", "Trace end-to-end inference layers.", "POST `/admin/api/trace` with `What is aPTT?`.", "Trace returns preprocessing, classifier, retrieval, and final route fields.", "Validated through admin trace runtime checks.", "Pass", "Manual + API"),
    TestCase("UT-22", "Unit", "Export pipeline", "Export accepted reviewed samples.", "GET `/admin/api/export-reviewed`.", "CSV file is generated for retraining.", "Validated through admin export flow.", "Pass", "Manual"),
    TestCase("UT-23", "Unit", "Export pipeline", "Export report-analysis failure candidates.", "GET `/admin/api/export-report-analysis-errors`.", "CSV file contains low-confidence / fallback report-analysis rows.", "Validated through admin export flow and store logic.", "Pass", "Manual + code review"),
    TestCase("UT-24", "Unit", "HTTPS deployment", "Build redirect URL correctly.", "Request `http://localhost:8000/admin?x=1` while HTTPS is enabled.", "Redirect points to HTTPS target preserving path and query.", "Passed in `test_https_redirect.py`.", "Pass", "Automated"),
    TestCase("UT-25", "Unit", "Certificate monitoring", "Classify near-expiry certificate severity.", "Evaluate certificate status with low remaining days.", "Admin warning level becomes warn/danger based on days remaining.", "Passed in `test_certificate_warning.py` and `test_ssl_utils.py`.", "Pass", "Automated"),
    TestCase("UT-26", "Unit", "Split pipeline", "Load fixed train/validation/test splits correctly.", "Read split metadata and split files.", "Split rows and ratios remain reproducible for model training.", "Passed in `test_split_utils.py`.", "Pass", "Automated"),
]


DETAILED_CASES = [
    ("UT-12", "Report numeric result analysis", FIG / "figure_3_3_chat_ui.png"),
    ("UT-15", "Pediatric report analysis", FIG / "figure_4_7_admin_trace.png"),
    ("UT-20", "Report-analysis preview API", FIG / "figure_4_5_admin_overview.png"),
    ("UT-24", "HTTPS redirect and deployment validation", FIG / "figure_4_5_admin_overview.png"),
]

AUTO_EVIDENCE = [
    (
        "Automated model-advisory regression evidence",
        "Figure A1 Automated pytest evidence covering UT-10 and UT-11.",
        FIG / "figure_t_auto_model_advisory.png",
    ),
    (
        "Automated report-analysis regression evidence",
        "Figure A2 Automated pytest evidence covering UT-12 to UT-18.",
        FIG / "figure_t_auto_report_analysis.png",
    ),
]


def set_default_font(style, name: str = "Arial", size: int = 11) -> None:
    style.font.name = name
    style.font.size = Pt(size)
    style._element.rPr.rFonts.set(qn("w:ascii"), name)
    style._element.rPr.rFonts.set(qn("w:hAnsi"), name)


def set_layout(doc: Document) -> None:
    for section in doc.sections:
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
    set_default_font(doc.styles["Normal"])
    set_default_font(doc.styles["Title"], size=20)
    set_default_font(doc.styles["Heading 1"], size=16)
    set_default_font(doc.styles["Heading 2"], size=13)


def set_spacing(paragraph, line_15: bool = True):
    fmt = paragraph.paragraph_format
    fmt.space_after = Pt(6)
    fmt.space_before = Pt(0)
    if line_15:
        fmt.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    else:
        fmt.line_spacing = 1.0


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def add_paragraph(doc: Document, text: str, style: str | None = None, single: bool = False):
    p = doc.add_paragraph(style=style)
    p.add_run(text)
    set_spacing(p, line_15=not single)
    return p


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.italic = True
    run.font.name = "Arial"
    run.font.size = Pt(10)
    set_spacing(p, line_15=False)


def add_image(doc: Document, path: Path, caption: str):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(5.6))
    set_spacing(p, line_15=False)
    add_caption(doc, caption)


def write_markdown() -> None:
    lines: list[str] = []
    lines.append("# Hybrid Chatbot Testing Documentation")
    lines.append("")
    lines.append("## 1. Test Plan")
    lines.append("")
    lines.append("This document records the unit-level and feature-level testing approach used for the Hybrid Chatbot for Medical Haematology. The scope covers the sequential intent classifier, report-analysis layer, retrieval-based routing, API behavior, admin monitoring flows, HTTPS deployment behavior, and retraining support exports.")
    lines.append("")
    lines.append("### 1.1 Objectives")
    lines.append("")
    lines.append("- Verify that the BiLSTM intent models classify supported hematology prompts correctly.")
    lines.append("- Verify that rule-based safety, entity detection, model auto-switching, and retrieval routing behave consistently.")
    lines.append("- Verify that the report model can analyze numeric CBC-style report phrases, sex-specific HGB/HCT phrases, and pediatric/age-aware phrases within bounded support rules.")
    lines.append("- Verify that the admin and MLOps support features export the correct retraining evidence.")
    lines.append("")
    lines.append("### 1.2 Environment")
    lines.append("")
    lines.append("- Operating environment: Windows 11 local development and Docker-based deployment")
    lines.append("- Backend: FastAPI")
    lines.append("- ML runtime: PyTorch BiLSTM intent classifier")
    lines.append("- Storage: PostgreSQL")
    lines.append("- Test framework: pytest")
    lines.append("- Latest automated suite status at documentation time: `50 passed`")
    lines.append("")
    lines.append("### 1.3 Entry and Exit Criteria")
    lines.append("")
    lines.append("- Entry criteria: features implemented, datasets merged, model artifacts generated, test environment available")
    lines.append("- Exit criteria: all planned cases executed, no blocking failures remaining, regression suite passing, evidence captured")
    lines.append("")
    lines.append("## 2. Test Case Catalogue")
    lines.append("")
    lines.append("| ID | Category | Component | Objective | Task | Expected Result | Actual Result | Status | Evidence |")
    lines.append("|---|---|---|---|---|---|---|---|---|")
    for tc in TEST_CASES:
        lines.append(f"| {tc.case_id} | {tc.category} | {tc.component} | {tc.objective} | {tc.task} | {tc.expected} | {tc.actual} | {tc.status} | {tc.evidence} |")
    lines.append("")
    lines.append("## 3. Automated Test Evidence")
    lines.append("")
    lines.append("The following figures provide screenshot-style evidence for the automated regression tests that covered the model-advisory and report-analysis cases listed in the catalogue.")
    lines.append("")
    lines.append("- Figure A1 covers the automated advisory checks mapped to UT-10 and UT-11.")
    lines.append("- Figure A2 covers the automated report-analysis and age-aware predictor checks mapped to UT-12 through UT-18.")
    lines.append("")
    lines.append("## 4. Detailed Test Execution Examples")
    lines.append("")
    for case_id, title, _ in DETAILED_CASES:
        tc = next(x for x in TEST_CASES if x.case_id == case_id)
        lines.append(f"### {case_id} {title}")
        lines.append("")
        lines.append(f"- Component: {tc.component}")
        lines.append(f"- Objective: {tc.objective}")
        lines.append(f"- Task: {tc.task}")
        lines.append(f"- Expected result: {tc.expected}")
        lines.append(f"- Actual result: {tc.actual}")
        lines.append(f"- Status: {tc.status}")
        lines.append("")
    lines.append("## 5. Execution Summary")
    lines.append("")
    lines.append("- Planned cases: 26")
    lines.append("- Passed: 26")
    lines.append("- Failed: 0")
    lines.append("- Blocked: 0")
    lines.append("- Automated regression suite at issue-close time: `50 passed`")
    lines.append("")
    lines.append("## 6. Residual Risks")
    lines.append("")
    lines.append("- Report analysis remains bounded and is not a diagnostic engine.")
    lines.append("- Reference intervals are generalized and must still defer to the printed report range.")
    lines.append("- Pediatric handling is deliberately simplified and should not be presented as full clinical decision support.")
    OUT_MD.write_text("\n".join(lines), encoding="utf-8")


def build_docx() -> None:
    doc = Document()
    set_layout(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("Hybrid Chatbot Testing Plan and Test Execution Documentation")
    r.bold = True
    r.font.name = "Arial"
    r.font.size = Pt(20)
    for line in [
        "",
        "Project: Hybrid Chatbot for Medical Haematology",
        "Document Type: Testing Documentation",
        "Prepared For: Dissertation Appendix / Supervisor Review",
        "Date: 03 May 2026",
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.add_run(line).font.name = "Arial"
        set_spacing(para, line_15=False)
    doc.add_page_break()

    add_paragraph(doc, "1. Test Plan", style="Heading 1")
    add_paragraph(doc, "This document recorded the unit-level and feature-level testing approach used for the Hybrid Chatbot for Medical Haematology. The scope covered the sequential intent classifier, report-analysis layer, retrieval-based routing, API behavior, admin monitoring flows, HTTPS deployment behavior, and retraining support exports.")
    add_paragraph(doc, "1.1 Objectives", style="Heading 2")
    for line in [
        "Verify that the BiLSTM intent models classified supported hematology prompts correctly.",
        "Verify that rule-based safety, entity detection, model auto-switching, and retrieval routing behaved consistently.",
        "Verify that the report model analyzed numeric CBC-style report phrases, sex-specific HGB/HCT phrases, and pediatric or age-aware phrases within bounded support rules.",
        "Verify that the admin and MLOps support features exported the correct retraining evidence.",
    ]:
        add_paragraph(doc, f"• {line}")
    add_paragraph(doc, "1.2 Test Environment", style="Heading 2")
    env_table = doc.add_table(rows=1, cols=2)
    env_table.style = "Table Grid"
    env_rows = [
        ("Operating environment", "Windows 11 local development and Docker-based deployment"),
        ("Backend", "FastAPI"),
        ("ML runtime", "PyTorch BiLSTM intent classifier"),
        ("Storage", "PostgreSQL"),
        ("Test framework", "pytest"),
        ("Latest automated suite", "50 passed"),
    ]
    env_table.rows[0].cells[0].text = "Item"
    env_table.rows[0].cells[1].text = "Value"
    for k, v in env_rows:
        cells = env_table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    for row in env_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_spacing(p, line_15=False)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10.5)

    add_paragraph(doc, "1.3 Entry and Exit Criteria", style="Heading 2")
    add_paragraph(doc, "Entry criteria: features implemented, datasets merged, model artifacts generated, and test environment available.")
    add_paragraph(doc, "Exit criteria: all planned cases executed, no blocking failures remaining, regression suite passing, and evidence captured.")

    add_paragraph(doc, "2. Test Case Catalogue", style="Heading 1")

    grouped = [
        ("2.1 Core inference and routing cases", TEST_CASES[0:10]),
        ("2.2 Report-analysis and prompt cases", TEST_CASES[10:18]),
        ("2.3 API, admin, export, and deployment cases", TEST_CASES[18:26]),
    ]
    for heading, cases in grouped:
        add_paragraph(doc, heading, style="Heading 2")
        cat_table = doc.add_table(rows=1, cols=7)
        cat_table.style = "Table Grid"
        headers = ["ID", "Component", "Objective", "Task", "Expected Result", "Actual Result", "Status / Evidence"]
        for cell, header in zip(cat_table.rows[0].cells, headers):
            cell.text = header
        for tc in cases:
            cells = cat_table.add_row().cells
            values = [
                tc.case_id,
                tc.component,
                tc.objective,
                tc.task,
                tc.expected,
                tc.actual,
                f"{tc.status} | {tc.evidence}",
            ]
            for cell, value in zip(cells, values):
                cell.text = value
        for row in cat_table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_spacing(p, line_15=False)
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(9.5)

    add_paragraph(doc, "3. Automated Test Evidence", style="Heading 1")
    add_paragraph(doc, "The following figures provide screenshot-style evidence for the automated regression tests that covered the advisory and report-analysis cases listed in the test catalogue.")
    for title, caption, image_path in AUTO_EVIDENCE:
        add_paragraph(doc, title, style="Heading 2")
        add_image(doc, image_path, caption)

    add_paragraph(doc, "4. Detailed Test Execution Examples", style="Heading 1")
    for idx, (case_id, title, image_path) in enumerate(DETAILED_CASES, start=1):
        tc = next(x for x in TEST_CASES if x.case_id == case_id)
        add_paragraph(doc, f"4.{idx} {case_id} {title}", style="Heading 2")
        meta = doc.add_table(rows=2, cols=3)
        meta.style = "Table Grid"
        meta.cell(0, 0).text = f"Test Case: {tc.case_id}"
        meta.cell(0, 1).text = f"Component: {tc.component}"
        meta.cell(0, 2).text = "Designed By: Project Author"
        meta.cell(1, 0).text = "Data Source: Code / API / Prompt Input"
        meta.cell(1, 1).text = f"Objective: {tc.objective}"
        meta.cell(1, 2).text = "Tester: Project Author"
        for row in meta.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    set_spacing(p, line_15=False)
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10.5)

        detail = doc.add_table(rows=2, cols=5)
        detail.style = "Table Grid"
        headers2 = ["Test Case", "Description", "Tasks", "Expected Result", "Actual Result"]
        for cell, header in zip(detail.rows[0].cells, headers2):
            cell.text = header
        row = detail.rows[1].cells
        row[0].text = tc.case_id
        row[1].text = tc.objective
        row[2].text = tc.task
        row[3].text = tc.expected
        row[4].text = f"{tc.actual} Status: {tc.status}."
        for r in detail.rows:
            for cell in r.cells:
                for p in cell.paragraphs:
                    set_spacing(p, line_15=False)
                    for run in p.runs:
                        run.font.name = "Arial"
                        run.font.size = Pt(10)
        add_image(doc, image_path, f"Figure T{idx} Supporting interface or evidence for {tc.case_id}")

    add_paragraph(doc, "5. Execution Summary", style="Heading 1")
    summary_table = doc.add_table(rows=1, cols=2)
    summary_table.style = "Table Grid"
    summary_table.rows[0].cells[0].text = "Metric"
    summary_table.rows[0].cells[1].text = "Value"
    for k, v in [
        ("Planned cases", "26"),
        ("Passed", "26"),
        ("Failed", "0"),
        ("Blocked", "0"),
        ("Automated regression suite", "50 passed"),
    ]:
        cells = summary_table.add_row().cells
        cells[0].text = k
        cells[1].text = v
    for row in summary_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                set_spacing(p, line_15=False)
                for run in p.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(10.5)

    add_paragraph(doc, "6. Residual Risks and Notes", style="Heading 1")
    add_paragraph(doc, "The report-analysis layer remained bounded and did not attempt diagnosis, treatment recommendation, or unrestricted clinical reasoning.")
    add_paragraph(doc, "Reference intervals remained generalized for assistant support and had to defer to the printed report range and local laboratory policy.")
    add_paragraph(doc, "Pediatric support remained deliberately simplified and should be documented as bounded laboratory assistance rather than clinical decision support.")

    for section in doc.sections:
        footer = section.footer.paragraphs[0]
        add_page_number(footer)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


def main() -> None:
    write_markdown()
    build_docx()
    print(OUT_MD)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
